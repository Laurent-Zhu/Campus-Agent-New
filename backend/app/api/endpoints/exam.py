from fastapi import APIRouter, Depends, HTTPException, Path, Response
from sqlalchemy.orm import Session
from typing import List
from ...core.deps import get_db, get_current_user
from ...schemas.exam import ExamCreate, Exam, ExamGenerateRequest, ExamUpdate
from ...models.exam import Exam as ExamModel, Question as QuestionModel
from ...models.user import User
from ai_agents.factory import AgentFactory
import datetime
import json
from fastapi.responses import JSONResponse
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from io import BytesIO
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from urllib.parse import quote

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm

from reportlab.lib.fonts import addMapping

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# from backend.app.main import vector_store
from ...core.deps import get_vector_store

pdfmetrics.registerFont(TTFont('SimHei', '/home/laurentzhu/PycharmProjects/CampusAgent-fusion/backend/app/static/fonts/simhei.ttf'))
addMapping('SimHei', 0, 0, 'SimHei')

router = APIRouter()

@router.post("/generate", response_model=ExamCreate)
async def generate_exam(
    request: ExamGenerateRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """生成考试题目"""
    print("backend/app/api/endpoints/exam.py的generate_exam正在生成考试试卷")
    if current_user.role != "teacher":
        raise HTTPException(
            status_code=403,
            detail="只有教师可以生成考试"
        )
    try:
        agent = AgentFactory.create_agent("exam_generator")
        if not agent:
            raise ValueError("创建智能体失败")
        print("generate_exam开始")
        vector_store = get_vector_store()
        questions_exam = await agent.generate_exam(
            course_id=request.course_id,
            knowledge_points=request.knowledge_points,
            question_config=request.question_types,  # 传递题型和数量配置
            difficulty=request.difficulty,
            duration=120,
            created_by=current_user.id,
            vector_store=vector_store  # 传递知识库
        )
        print("generate_exam成功")
        return questions_exam
    except Exception as e:
        print("生成考试异常:", e)
        raise HTTPException(
            status_code=500,
            detail=f"生成考试失败: {str(e)}"
        )

def generate_word_from_exam_data(exam_data: dict, include_analysis: bool = True) -> BytesIO:
    document = Document()

    # 设置默认字体为 SimHei（黑体）
    style = document.styles['Normal']
    font = style.font
    font.name = 'SimHei'
    font.size = Pt(12)
    # 设置中文字体（Windows 下有效，Linux 下无影响）
    r = style.element.rPr
    r.rFonts.set(qn('w:eastAsia'), 'SimHei')

    # 标题
    document.add_heading(exam_data.get("title", "试卷"), level=1)

    # 题目
    for idx, q in enumerate(exam_data.get("questions", []), 1):
        document.add_paragraph(f"Q{idx}: {q.get('content', '')}")

        options = q.get("options", [])
        if isinstance(options, list):
            for i, opt in enumerate(options):
                label = chr(65 + i)  # A, B, C...
                if isinstance(opt, dict):
                    text = f"{label}. {opt.get('text', '')}"
                else:
                    text = f"{label}. {str(opt)}"
                document.add_paragraph(text, style='List Bullet')

        # 根据 include_analysis 参数决定是否包含答案和解析
        if include_analysis:
            document.add_paragraph(f"答案: {q.get('answer', '')}")
            document.add_paragraph(f"解析: {q.get('analysis', '')}")
        document.add_paragraph("")  # 空行

    # 保存到内存 buffer
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def generate_pdf_from_exam_data(exam_data: dict, include_analysis: bool = True) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=20*mm, rightMargin=20*mm, topMargin=20*mm, bottomMargin=20*mm)
    styles = getSampleStyleSheet()
    styles['Normal'].fontName = 'SimHei'
    styles['Title'].fontName = 'SimHei'

    story = []
    story.append(Paragraph(exam_data.get("title", "试卷"), styles['Title']))
    story.append(Spacer(1, 12))

    for idx, q in enumerate(exam_data.get("questions", []), 1):
        question_text = f"Q{idx}: {q.get('content', '')}"
        story.append(Paragraph(question_text, styles['Normal']))
        story.append(Spacer(1, 6))

        options = q.get("options", [])
        if isinstance(options, list):
            for opt in options:
                story.append(Paragraph(opt, styles['Normal']))
                story.append(Spacer(1, 3))

        if include_analysis:
            story.append(Paragraph(f"答案: {q.get('answer', '')}", styles['Normal']))
            story.append(Spacer(1, 3))
            story.append(Paragraph(f"解析: {q.get('analysis', '')}", styles['Normal']))
            story.append(Spacer(1, 12))

    doc.build(story)
    buffer.seek(0)
    return buffer

@router.post("/generate-pdf")
async def generate_exam_pdf(exam_data: dict, include_analysis: bool = True):
    """
    接收前端传来的试卷内容，生成PDF并返回文件流
    """
    try:
        pdf_buffer = generate_pdf_from_exam_data(exam_data, include_analysis)
        filename = f"{exam_data.get('title', 'exam')}.pdf"
        encoded_filename = quote(filename)

        return Response(
            pdf_buffer.read(),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
            }
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"PDF生成失败: {str(e)}")

@router.post("/generate-word")
async def generate_exam_word(exam_data: dict, include_analysis: bool = True):
    """
    接收前端传来的试卷内容，生成Word文档并返回文件流
    """
    try:
        word_buffer = generate_word_from_exam_data(exam_data, include_analysis)
        filename = f"{exam_data.get('title', 'exam')}.docx"
        encoded_filename = quote(filename)

        return Response(
            content=word_buffer.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
            }
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Word生成失败: {str(e)}")

async def create_exam(
    exam: ExamCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """保存考试"""
    if current_user.role != "teacher":
        raise HTTPException(
            status_code=403,
            detail="只有教师可以创建考试"
        )
    try:
        db_exam = ExamModel(
            title=exam.title,
            description=exam.description,
            course_id=exam.course_id,
            duration=exam.duration,
            total_score=exam.total_score,
            created_by=current_user.id,
            created_at=datetime.datetime.now(),
            status="draft"
        )
        db.add(db_exam)
        db.commit()
        db.refresh(db_exam)
        # 添加题目
        for q in exam.questions:
            db_question = QuestionModel(
                type=q.type,
                content=q.content,
                options=json.dumps(q.options) if q.options else None,
                answer=q.answer,
                analysis=q.analysis,
                score=q.score,
                knowledge_point=q.knowledge_point,
                difficulty=q.difficulty,
                exam_id=db_exam.id
            )
            db.add(db_question)
        db.commit()
        db.refresh(db_exam)
        return db_exam
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=500,
            detail=f"创建考试失败: {str(e)}"
        )

@router.get("/{exam_id}", response_model=Exam)
async def get_exam(
    exam_id: int = Path(..., title="考试ID"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取考试详情"""
    exam = db.query(ExamModel).filter(ExamModel.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="考试不存在")
    
    if exam.created_by != current_user.id and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="无权访问此考试")
    
    return exam

@router.put("/{exam_id}", response_model=Exam)
async def update_exam(
    exam_id: int,
    exam_update: ExamUpdate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """更新考试"""
    db_exam = db.query(ExamModel).filter(ExamModel.id == exam_id).first()
    if not db_exam:
        raise HTTPException(status_code=404, detail="考试不存在")
    
    if db_exam.created_by != current_user.id:
        raise HTTPException(status_code=403, detail="只有创建者可以修改考试")
    
    try:
        for key, value in exam_update.dict(exclude_unset=True).items():
            setattr(db_exam, key, value)
        
        db.commit()
        db.refresh(db_exam)
        return db_exam
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=500,
            detail=f"更新考试失败: {str(e)}"
        )

@router.delete("/{exam_id}")
async def delete_exam(
    exam_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """删除考试"""
    db_exam = db.query(ExamModel).filter(ExamModel.id == exam_id).first()
    if not db_exam:
        raise HTTPException(status_code=404, detail="考试不存在")
    
    if db_exam.created_by != current_user.id:
        raise HTTPException(status_code=403, detail="只有创建者可以删除考试")
    
    try:
        db.delete(db_exam)
        db.commit()
        return {"message": "考试已删除"}
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=500,
            detail=f"删除考试失败: {str(e)}"
        )

@router.get("/download/{exam_id}")
async def download_exam(
    exam_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """下载试卷（JSON格式，可扩展为Word/PDF）"""
    exam = db.query(ExamModel).filter(ExamModel.id == exam_id).first()
    if not exam:
        raise HTTPException(status_code=404, detail="试卷不存在")
    # 权限校验
    if exam.created_by != current_user.id and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="无权下载此试卷")
    # 组装题目
    questions = db.query(QuestionModel).filter(QuestionModel.exam_id == exam_id).all()
    exam_dict = {
        "id": exam.id,
        "title": exam.title,
        "description": exam.description,
        "course_id": exam.course_id,
        "duration": exam.duration,
        "total_score": exam.total_score,
        "created_at": exam.created_at,
        "created_by": exam.created_by,
        "status": exam.status,
        "questions": [
            {
                "id": q.id,
                "type": q.type,
                "content": q.content,
                "options": json.loads(q.options) if q.options else None,
                "answer": q.answer,
                "analysis": q.analysis,
                "score": q.score,
                "knowledge_point": q.knowledge_point,
                "difficulty": q.difficulty
            } for q in questions
        ]
    }
    return JSONResponse(content=exam_dict)